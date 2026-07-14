"""Multimodal document parsing for BookRAG.

Turns any corpus file into a flat list of `Element`s (narrative text or a table),
each carrying page/section provenance. Tables are serialised to Markdown so the
embedder and LLM can actually reason over the numbers -- the whole point of going
multimodal for a finance corpus.

Routing:
  .pdf            -> PyMuPDF text per page + native table detection; OCR fallback
                     for image-only pages (only if an OCR engine is available).
  .htm/.html      -> BeautifulSoup narrative + financial tables (SEC filings).

Design principles: never let one bad page/table/file crash a run; degrade gracefully.
"""
from __future__ import annotations

import io
import warnings
from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from typing import Iterable, Optional

import config

warnings.filterwarnings("ignore")

# Optional cap on pages parsed per PDF (set by ingest.py --max-pages for fast tests).
MAX_PAGES: Optional[int] = None


@dataclass
class Element:
    text: str
    element_type: str = "text"        # "text" | "table"
    page: Optional[int] = None        # 1-indexed page (PDF) or section index (HTML)
    extra: dict = field(default_factory=dict)


# ------------------------------------------------------------------- OCR (optional)
@lru_cache(maxsize=1)
def _ocr_engine():
    """Return a callable(image_bytes)->str, or None if no OCR engine is installed.

    Tries pytesseract (needs `brew install tesseract`) then rapidocr (pip-only).
    Controlled by config.ENABLE_OCR ("auto"|"on"|"off").
    """
    if config.ENABLE_OCR == "off":
        return None
    try:
        import pytesseract  # noqa: F401
        from PIL import Image

        def run(png_bytes: bytes) -> str:
            return pytesseract.image_to_string(Image.open(io.BytesIO(png_bytes)))

        return run
    except Exception:
        pass
    try:
        from rapidocr_onnxruntime import RapidOCR

        engine = RapidOCR()

        def run(png_bytes: bytes) -> str:  # type: ignore[misc]
            result, _ = engine(png_bytes)
            return "\n".join(line[1] for line in (result or []))

        return run
    except Exception:
        return None


# ------------------------------------------------------------------- helpers
def _df_to_markdown(df) -> str:
    """DataFrame -> GitHub Markdown table without requiring `tabulate`.

    If the columns are just an integer RangeIndex (no real header row, common in
    SEC filings), the header line is dropped so the LLM isn't fed "| 0 | 1 |" noise.
    """
    df = df.fillna("")
    has_header = not all(isinstance(c, int) for c in df.columns)
    ncols = df.shape[1]
    rows = [
        "| " + " | ".join(str(v).replace("\n", " ").strip() for v in row) + " |"
        for row in df.itertuples(index=False, name=None)
    ]
    if has_header:
        cols = [str(c).replace("\n", " ").strip() for c in df.columns]
        header = "| " + " | ".join(cols) + " |"
        sep = "| " + " | ".join("---" for _ in cols) + " |"
        return "\n".join([header, sep, *rows])
    sep = "| " + " | ".join("---" for _ in range(ncols)) + " |"
    return "\n".join([sep, *rows])


def _clean_df(df):
    """Drop empty + near-empty rows/cols (SEC colspan spacers) and normalise cells."""
    df = df.replace(r"^\s*$", None, regex=True)
    df = df.dropna(axis=1, how="all").dropna(axis=0, how="all")
    # Colspan expansion leaves spacer columns/rows with a single spanned value;
    # drop anything with fewer than 2 real cells once the table is big enough.
    if df.shape[1] > 2:
        df = df.loc[:, df.notna().sum(axis=0) >= 2]
    if df.shape[0] > 2:
        df = df.loc[df.notna().sum(axis=1) >= 2]
    return df.fillna("")


def _looks_like_data_table(df) -> bool:
    """Keep real financial/data tables; drop layout & single-value spacer tables."""
    try:
        rows, cols = df.shape
    except Exception:
        return False
    if rows < 2 or cols < 2 or rows * cols < config.MIN_TABLE_CELLS:
        return False
    cells = [str(v) for v in df.to_numpy().ravel() if str(v).strip()]
    if not cells:
        return False
    joined = " ".join(cells)
    numeric = sum(1 for c in cells if any(ch.isdigit() for ch in c))
    # Financial tables are number-dense or carry currency/percent markers.
    return (numeric / len(cells)) >= 0.15 or "$" in joined or "%" in joined


# ------------------------------------------------------------------- PDF
def parse_pdf(path: Path) -> list[Element]:
    import fitz

    elements: list[Element] = []
    try:
        doc = fitz.open(path)
    except Exception as e:  # noqa: BLE001
        print(f"    ! cannot open PDF {path.name}: {e}")
        return elements

    ocr = _ocr_engine()
    n_pages = doc.page_count if MAX_PAGES is None else min(doc.page_count, MAX_PAGES)
    for pno in range(n_pages):
        try:
            page = doc[pno]
            text = page.get_text("text").strip()

            if len(text) < config.MIN_CHARS_PER_PAGE and ocr is not None:
                try:
                    pix = page.get_pixmap(dpi=200)
                    ocr_text = ocr(pix.tobytes("png")).strip()
                    if len(ocr_text) > len(text):
                        text = ocr_text
                except Exception:
                    pass

            if text:
                elements.append(Element(text=text, element_type="text", page=pno + 1))

            # Native table extraction (fast, no extra dependency).
            try:
                for tbl in page.find_tables().tables:
                    md = tbl.to_markdown().replace("<br>", " ").strip()
                    if md and md.count("|") >= config.MIN_TABLE_CELLS:
                        elements.append(
                            Element(text=md, element_type="table", page=pno + 1)
                        )
            except Exception:
                pass
        except Exception as e:  # noqa: BLE001 - skip the page, keep the doc
            print(f"    ~ page {pno + 1} of {path.name} skipped: {e}")
            continue

    doc.close()
    return elements


# ------------------------------------------------------------------- HTML (SEC filings)
def parse_html(path: Path) -> list[Element]:
    import pandas as pd
    from bs4 import BeautifulSoup

    raw = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(raw, "lxml")
    for tag in soup(["script", "style", "head", "title"]):
        tag.decompose()

    elements: list[Element] = []

    # 1) Financial tables -> Markdown. SEC filings are XHTML, so parse each <table>
    #    tag individually with lxml (a whole-document read_html would demand html5lib).
    kept = 0
    for tag in soup.find_all("table"):
        try:
            df = pd.read_html(io.StringIO(str(tag)), flavor="lxml")[0]
            df = _clean_df(df)
        except Exception:
            continue
        if _looks_like_data_table(df):
            elements.append(
                Element(text=_df_to_markdown(df), element_type="table",
                        page=None, extra={"table_index": kept})
            )
            kept += 1

    # 2) Narrative text with tables removed (so numbers aren't duplicated as prose soup).
    for t in soup.find_all("table"):
        t.decompose()
    narrative = soup.get_text(separator="\n")
    lines = [ln.strip() for ln in narrative.splitlines()]
    narrative = "\n".join(ln for ln in lines if ln)
    if narrative:
        elements.append(Element(text=narrative, element_type="text", page=None))

    return elements


# ------------------------------------------------------------------- images (OCR)
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"}


def parse_image(path: Path) -> list[Element]:
    """OCR an uploaded image into searchable text. Graceful if no OCR engine."""
    ocr = _ocr_engine()
    if ocr is None:
        return [Element(
            text=f"[Image: {path.name}] — no OCR engine installed; install pytesseract "
                 "(brew install tesseract) or rapidocr-onnxruntime to extract text.",
            element_type="image", extra={"filename": path.name})]
    try:
        text = ocr(path.read_bytes()).strip()
    except Exception as e:  # noqa: BLE001
        text = ""
        print(f"    ! OCR failed for {path.name}: {e}")
    if not text:
        text = f"[Image: {path.name}] — no text detected."
    return [Element(text=text, element_type="image", extra={"filename": path.name})]


# ------------------------------------------------------------------- audio (transcription)
AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".flac", ".ogg", ".aac", ".webm", ".mp4"}


@lru_cache(maxsize=1)
def _transcriber():
    """Return a callable(path)->str transcript, or None if no engine is installed.

    Tries faster-whisper (fast, CPU int8) then openai-whisper. Controlled by
    config.ENABLE_AUDIO ("auto"|"off").
    """
    if getattr(config, "ENABLE_AUDIO", "auto") == "off":
        return None
    model_size = getattr(config, "WHISPER_MODEL", "base")
    try:
        from faster_whisper import WhisperModel
        model = WhisperModel(model_size, device="cpu", compute_type="int8")

        def run(path) -> str:
            segments, _ = model.transcribe(str(path))
            return " ".join(seg.text for seg in segments)

        return run
    except Exception:
        pass
    try:
        import whisper
        model = whisper.load_model(model_size)

        def run(path) -> str:  # type: ignore[misc]
            return model.transcribe(str(path)).get("text", "")

        return run
    except Exception:
        return None


def parse_audio(path: Path) -> list[Element]:
    """Transcribe an uploaded audio/video file into searchable text. Graceful if no engine."""
    tr = _transcriber()
    if tr is None:
        return [Element(
            text=f"[Audio: {path.name}] — no transcription engine installed; "
                 "install faster-whisper to transcribe.",
            element_type="audio", extra={"filename": path.name})]
    try:
        text = tr(path).strip()
    except Exception as e:  # noqa: BLE001
        text = ""
        print(f"    ! transcription failed for {path.name}: {e}")
    if not text:
        text = f"[Audio: {path.name}] — no speech detected."
    return [Element(text=text, element_type="audio", extra={"filename": path.name})]


# ------------------------------------------------------------------- docx
def parse_docx(path: Path) -> list[Element]:
    try:
        import docx  # python-docx
    except Exception:
        return [Element(text=f"[Document: {path.name}] — install python-docx to extract text.",
                        element_type="text", extra={"filename": path.name})]
    d = docx.Document(str(path))
    text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
    return [Element(text=text, element_type="text")] if text else []


# ------------------------------------------------------------------- dispatcher
def parse_document(path: Path) -> list[Element]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return parse_pdf(path)
    if ext in (".htm", ".html"):
        return parse_html(path)
    if ext in (".txt", ".md", ".csv"):
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        return [Element(text=text, element_type="text")] if text else []
    if ext == ".docx":
        return parse_docx(path)
    if ext in IMAGE_EXTS:
        return parse_image(path)
    if ext in AUDIO_EXTS:
        return parse_audio(path)
    print(f"    ? unsupported type {ext} ({path.name}) -- skipped")
    return []


def ocr_available() -> bool:
    return _ocr_engine() is not None


def audio_available() -> bool:
    return _transcriber() is not None
